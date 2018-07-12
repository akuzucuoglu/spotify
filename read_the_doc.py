# -*- coding: utf-8 -*-
"""
Created on Tue Jul 10 10:10:00 2018

@author: ahmet.kuzucuoglu
"""

print('Ja')

#1.Doküman okumak için gereken library import edilir.
import docx
#2. String to date operasyonları için gereken library import edilir.
from datetime import datetime

#Document >> Paragraphs >> Runs (ve .text)
#Document >> Tables >> Rows || Columns >> Cells 

#2. Benim oluşturduğum kopya üzerinde yaptıgım çalışma
#my_doc= docx.Document('Yerel_Yonetici_Yetki_Formu.docx')
my_doc= docx.Document('Uzak_Masaustu_Yetki_Formu.docx')



#3. Temel şeyler
len(my_doc.paragraphs)  #Kaç tane paragraph objesi olduğunu gösterir.
len(my_doc.tables)      #Kaç tane table objesi olduğunu gösterir.

#4. Table üzerinde çalışmak için
# Dokümandaki ilk tablo alınır.
my_table= my_doc.tables[0]

#deneme işlemleri -- Burada kontrolleri yaptırsan tatlış olur.
my_table.rows[0].cells[1].text


my_dict={} #Dictionary yaratılır
my_list=[] #List yaratılır
my_tuple=() #Tuple yaratılır
my_RowList=[] #Insert edilecek rowList


#Tablodaki her bir hücreyi sırayla döndüm
for row in my_table.rows:
    if row.cells[0].text=='Yetki Baslangıc Tarihi:' or row.cells[0].text=='Yetki Bitis Tarihi:':
        #my_dict[row.cells[0].text]=datetime.strptime(row.cells[1].text,'%d.%m.%Y')
        my_dict[row.cells[0].text]=row.cells[1].text
        my_list.append(datetime.strptime(row.cells[1].text,'%d.%m.%Y'))
    else:
        my_dict[row.cells[0].text]=row.cells[1].text
        my_list.append(row.cells[1].text)

print(my_dict)

#List objesini veritabanına insert edilebilir bir hale getirebilmek için öncetuple objesine dönüştürülür.  
#Sonra bu tuple yeni bir list objesine assign edilir.      
my_tuple=tuple(my_list)   
my_RowList.append(my_tuple)

print(my_RowList)

#Veritabanına yazdırma işlemi gerçekleştirilir.
import cx_Oracle
dsn_tns = cx_Oracle.makedsn('IDB-10256.roketsan.com.tr', '1521', service_name='XE')
conn = cx_Oracle.connect(user='HR', password='HR', dsn=dsn_tns)

import sys
sys.getdefaultencoding()

my_RowList



cur_insert = conn.cursor()
cur_insert.bindarraysize = 1
cur_insert.setinputsizes(int, 4000)
cur_insert.executemany("INSERT INTO xxrks_sdi_grants2(form_name,grantee,granter,grant_start_date,grant_end_date,reason,result) VALUES(:form_name,:grantee,:3,:grant_start_date,:grant_end_date,:reason,:result)", my_RowList)
cur_insert.close()
conn.commit()

# Now query the results back
cur_select = conn.cursor()
cur_select.execute('select * from xxrks_sdi_grants')
res = cur_select.fetchall()
print(res)               
cur_select.close()



#Oluşturduğum dictionary'i de yine bir json dosyasına yazdırayım
import json
with open('my_dict.json', 'w') as outfile:
    json.dump(my_dict, outfile)
    

#Oluşturduğum json'ı tekrar okuyalım.    
from pprint import pprint

with open('my_dict.json') as f:
    data = json.load(f)

pprint(data)




print('ja')